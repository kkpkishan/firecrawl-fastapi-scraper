"""
Document Content Extraction Module

Extracts text and data from various document formats:
- PDF (.pdf)
- Excel (.xlsx, .xls)
- Word (.docx, .doc)
- OpenDocument (.odt)
- Rich Text (.rtf)
"""
import logging
import httpx
import tempfile
import os
from typing import Optional, Dict, Tuple
from pathlib import Path

logger = logging.getLogger(__name__)


class DocumentExtractor:
    """
    Extract content from various document formats.
    """
    
    # Supported document extensions
    SUPPORTED_EXTENSIONS = {
        '.pdf', '.xlsx', '.xls', '.docx', '.doc', '.odt', '.rtf'
    }
    
    def __init__(self):
        """Initialize document extractor."""
        self.timeout = 60.0  # 60 seconds for document downloads
        logger.info(f"DocumentExtractor initialized. Supported formats: {self.SUPPORTED_EXTENSIONS}")
    
    def is_document_url(self, url: str) -> bool:
        """
        Check if URL points to a supported document.
        
        Args:
            url: URL to check
            
        Returns:
            True if URL is a document, False otherwise
        """
        url_lower = url.lower()
        return any(url_lower.endswith(ext) for ext in self.SUPPORTED_EXTENSIONS)
    
    def get_document_type(self, url: str) -> Optional[str]:
        """
        Get document type from URL.
        
        Args:
            url: Document URL
            
        Returns:
            Document extension (e.g., '.pdf') or None
        """
        url_lower = url.lower()
        for ext in self.SUPPORTED_EXTENSIONS:
            if url_lower.endswith(ext):
                return ext
        return None
    
    async def download_document(self, url: str) -> Tuple[Optional[bytes], Optional[str]]:
        """
        Download document from URL.
        
        Args:
            url: Document URL
            
        Returns:
            Tuple of (document_bytes, error_message)
        """
        try:
            logger.info(f"Downloading document: {url}")
            
            async with httpx.AsyncClient(timeout=self.timeout, follow_redirects=True) as client:
                response = await client.get(url)
                
                if response.status_code == 200:
                    logger.info(f"Successfully downloaded document ({len(response.content)} bytes)")
                    return response.content, None
                else:
                    error_msg = f"Failed to download document: HTTP {response.status_code}"
                    logger.error(error_msg)
                    return None, error_msg
                    
        except httpx.TimeoutException:
            error_msg = f"Timeout downloading document from {url}"
            logger.error(error_msg)
            return None, error_msg
        except Exception as e:
            error_msg = f"Error downloading document: {str(e)}"
            logger.error(error_msg, exc_info=True)
            return None, error_msg
    
    async def extract_text(self, url: str) -> Tuple[Optional[str], Optional[str]]:
        """
        Extract text content from document URL.
        
        Args:
            url: Document URL
            
        Returns:
            Tuple of (extracted_text, error_message)
        """
        # Check if supported document
        doc_type = self.get_document_type(url)
        if not doc_type:
            return None, f"Unsupported document type: {url}"
        
        # Download document
        doc_bytes, error = await self.download_document(url)
        if not doc_bytes:
            return None, error
        
        # Extract based on document type
        try:
            if doc_type == '.pdf':
                return await self._extract_pdf(doc_bytes)
            elif doc_type in ['.xlsx', '.xls']:
                return await self._extract_excel(doc_bytes)
            elif doc_type in ['.docx', '.doc']:
                return await self._extract_word(doc_bytes)
            elif doc_type == '.odt':
                return await self._extract_odt(doc_bytes)
            elif doc_type == '.rtf':
                return await self._extract_rtf(doc_bytes)
            else:
                return None, f"Extraction not implemented for {doc_type}"
                
        except Exception as e:
            error_msg = f"Error extracting {doc_type}: {str(e)}"
            logger.error(error_msg, exc_info=True)
            return None, error_msg
    
    async def extract_metadata(self, url: str) -> Dict[str, any]:
        """
        Extract metadata from document URL.
        
        Args:
            url: Document URL
            
        Returns:
            Dictionary with document metadata
        """
        metadata = {
            'document_type': None,
            'page_count': None,
            'sheet_names': None,
            'section_count': None,
            'error': None
        }
        
        # Check if supported document
        doc_type = self.get_document_type(url)
        if not doc_type:
            metadata['error'] = f"Unsupported document type: {url}"
            return metadata
        
        metadata['document_type'] = doc_type
        
        # Download document
        doc_bytes, error = await self.download_document(url)
        if not doc_bytes:
            metadata['error'] = error
            return metadata
        
        # Extract metadata based on document type
        try:
            if doc_type == '.pdf':
                metadata.update(await self._extract_pdf_metadata(doc_bytes))
            elif doc_type in ['.xlsx', '.xls']:
                metadata.update(await self._extract_excel_metadata(doc_bytes))
            elif doc_type in ['.docx', '.doc']:
                metadata.update(await self._extract_word_metadata(doc_bytes))
        except Exception as e:
            metadata['error'] = f"Error extracting metadata: {str(e)}"
            logger.error(f"Metadata extraction error for {doc_type}: {e}", exc_info=True)
        
        return metadata
    
    async def _extract_pdf_metadata(self, pdf_bytes: bytes) -> Dict[str, any]:
        """Extract metadata from PDF."""
        try:
            import PyPDF2
            import io
            
            pdf_file = io.BytesIO(pdf_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            return {
                'page_count': len(pdf_reader.pages)
            }
        except Exception as e:
            logger.warning(f"Error extracting PDF metadata: {e}")
            return {}
    
    async def _extract_excel_metadata(self, excel_bytes: bytes) -> Dict[str, any]:
        """Extract metadata from Excel."""
        try:
            import openpyxl
            import io
            
            excel_file = io.BytesIO(excel_bytes)
            workbook = openpyxl.load_workbook(excel_file, data_only=True)
            
            return {
                'sheet_names': workbook.sheetnames
            }
        except Exception as e:
            logger.warning(f"Error extracting Excel metadata: {e}")
            return {}
    
    async def _extract_word_metadata(self, word_bytes: bytes) -> Dict[str, any]:
        """Extract metadata from Word document."""
        try:
            import docx
            import io
            
            word_file = io.BytesIO(word_bytes)
            doc = docx.Document(word_file)
            
            # Count sections (paragraphs with heading styles)
            section_count = 0
            for para in doc.paragraphs:
                if para.style.name.startswith('Heading'):
                    section_count += 1
            
            return {
                'section_count': section_count if section_count > 0 else len(doc.paragraphs)
            }
        except Exception as e:
            logger.warning(f"Error extracting Word metadata: {e}")
            return {}
    
    async def _extract_pdf(self, pdf_bytes: bytes) -> Tuple[Optional[str], Optional[str]]:
        """Extract text from PDF."""
        try:
            import PyPDF2
            import io
            
            pdf_file = io.BytesIO(pdf_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text:
                        text_parts.append(f"--- Page {page_num + 1} ---\n{text}")
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num + 1}: {e}")
                    continue
            
            if text_parts:
                full_text = "\n\n".join(text_parts)
                logger.info(f"Extracted {len(full_text)} characters from PDF ({len(pdf_reader.pages)} pages)")
                return full_text, None
            else:
                return None, "No text extracted from PDF"
                
        except ImportError:
            return None, "PyPDF2 library not installed"
        except Exception as e:
            return None, f"PDF extraction error: {str(e)}"
    
    async def _extract_excel(self, excel_bytes: bytes) -> Tuple[Optional[str], Optional[str]]:
        """Extract text from Excel."""
        try:
            import openpyxl
            import io
            
            excel_file = io.BytesIO(excel_bytes)
            workbook = openpyxl.load_workbook(excel_file, data_only=True)
            
            text_parts = []
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"--- Sheet: {sheet_name} ---")
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = "\t".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        text_parts.append(row_text)
            
            if text_parts:
                full_text = "\n".join(text_parts)
                logger.info(f"Extracted {len(full_text)} characters from Excel ({len(workbook.sheetnames)} sheets)")
                return full_text, None
            else:
                return None, "No data extracted from Excel"
                
        except ImportError:
            return None, "openpyxl library not installed"
        except Exception as e:
            return None, f"Excel extraction error: {str(e)}"
    
    async def _extract_word(self, word_bytes: bytes) -> Tuple[Optional[str], Optional[str]]:
        """Extract text from Word document."""
        try:
            import docx
            import io
            
            word_file = io.BytesIO(word_bytes)
            doc = docx.Document(word_file)
            
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                text_parts.append("\n--- Table ---")
                for row in table.rows:
                    row_text = "\t".join([cell.text for cell in row.cells])
                    text_parts.append(row_text)
            
            if text_parts:
                full_text = "\n".join(text_parts)
                logger.info(f"Extracted {len(full_text)} characters from Word document")
                return full_text, None
            else:
                return None, "No text extracted from Word document"
                
        except ImportError:
            return None, "python-docx library not installed"
        except Exception as e:
            return None, f"Word extraction error: {str(e)}"
    
    async def _extract_odt(self, odt_bytes: bytes) -> Tuple[Optional[str], Optional[str]]:
        """Extract text from OpenDocument."""
        try:
            from odf import text, teletype
            from odf.opendocument import load
            import io
            
            # Save to temp file (odfpy requires file path)
            with tempfile.NamedTemporaryFile(delete=False, suffix='.odt') as tmp:
                tmp.write(odt_bytes)
                tmp_path = tmp.name
            
            try:
                doc = load(tmp_path)
                text_parts = []
                
                for para in doc.getElementsByType(text.P):
                    para_text = teletype.extractText(para)
                    if para_text.strip():
                        text_parts.append(para_text)
                
                if text_parts:
                    full_text = "\n".join(text_parts)
                    logger.info(f"Extracted {len(full_text)} characters from ODT")
                    return full_text, None
                else:
                    return None, "No text extracted from ODT"
            finally:
                # Clean up temp file
                os.unlink(tmp_path)
                
        except ImportError:
            return None, "odfpy library not installed"
        except Exception as e:
            return None, f"ODT extraction error: {str(e)}"
    
    async def _extract_rtf(self, rtf_bytes: bytes) -> Tuple[Optional[str], Optional[str]]:
        """Extract text from RTF."""
        try:
            from striprtf.striprtf import rtf_to_text
            
            rtf_text = rtf_bytes.decode('utf-8', errors='ignore')
            plain_text = rtf_to_text(rtf_text)
            
            if plain_text and plain_text.strip():
                logger.info(f"Extracted {len(plain_text)} characters from RTF")
                return plain_text, None
            else:
                return None, "No text extracted from RTF"
                
        except ImportError:
            return None, "striprtf library not installed"
        except Exception as e:
            return None, f"RTF extraction error: {str(e)}"


# Global extractor instance
_document_extractor_instance = None


def get_document_extractor() -> DocumentExtractor:
    """
    Get global DocumentExtractor instance (singleton).
    
    Returns:
        DocumentExtractor instance
    """
    global _document_extractor_instance
    if _document_extractor_instance is None:
        _document_extractor_instance = DocumentExtractor()
    return _document_extractor_instance
